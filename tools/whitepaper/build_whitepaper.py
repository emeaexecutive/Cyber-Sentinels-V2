#!/usr/bin/env python3
"""Build the designed Cyber Sentinels V1 whitepaper from its canonical Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "whitepaper" / "CYBER_SENTINELS_WHITEPAPER_V1.md"
OUTPUT = ROOT / "tmp" / "whitepaper" / "CYBER_SENTINELS_WHITEPAPER_V1.docx"

INK = "081018"
CHARCOAL = "111B24"
CYAN = "11D9E8"
PALE_CYAN = "DDFBFF"
SOFT_GREY = "647481"
LIGHT = "EFF5F6"
WHITE = "FFFFFF"
GREEN = "0B866E"
AMBER = "B96E00"
RED = "A63A43"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc.get_or_add_tcPr()
    node = tc.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        tc.append(node)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        part = node.find(qn(f"w:{name}"))
        if part is None:
            part = OxmlElement(f"w:{name}")
            node.append(part)
        part.set(qn("w:w"), str(value))
        part.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, color=INK, size=9, bold=False, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def borderless(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_rule(doc, color=CYAN, height="18") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), height)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("CYBER SENTINELS  ·  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(SOFT_GREY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name, size, before, after, color in (
        ("Title", 30, 0, 10, WHITE),
        ("Heading 1", 20, 0, 7, INK),
        ("Heading 2", 13, 9, 4, CHARCOAL),
        ("Heading 3", 11, 7, 3, CHARCOAL),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    props = doc.core_properties
    props.title = "Cyber Sentinels: Operational Trust Infrastructure for Autonomous Systems"
    props.subject = "Technical Whitepaper, Version 1.0"
    props.author = "Cyber Sentinels"
    props.keywords = "operational trust, AI agents, authority, evidence, receipt, Replay"
    props.comments = "Public technical whitepaper."


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.9)
    borderless(table)
    cell = table.cell(0, 0)
    shade(cell, INK)
    margins(cell, 700, 520, 650, 520)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(38)
    r = p.add_run("CYBER SENTINELS")
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(CYAN)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Operational Trust\nInfrastructure for\nAutonomous Systems")
    r.font.name = "Calibri"
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("The control layer between autonomous intelligence and real-world action.")
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(PALE_CYAN)

    p = cell.add_paragraph()
    r = p.add_run("TECHNICAL WHITEPAPER\nVERSION 1.0  ·  2026")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(CYAN)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("www.cybersentinels.com")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(WHITE)


def page_header(doc: Document, number: str, title: str, kicker: str | None = None) -> None:
    row = doc.add_table(rows=1, cols=2)
    row.autofit = False
    row.columns[0].width = Inches(0.78)
    row.columns[1].width = Inches(5.85)
    row.alignment = WD_TABLE_ALIGNMENT.CENTER
    borderless(row)
    shade(row.cell(0, 0), CYAN)
    shade(row.cell(0, 1), INK)
    set_cell_text(row.cell(0, 0), number, color=INK, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cell(0, 1), kicker or "OPERATIONAL TRUST INFRASTRUCTURE", color=WHITE, size=8, bold=True)
    p = doc.add_paragraph(title, style="Heading 1")
    p.paragraph_format.space_before = Pt(15)
    add_rule(doc)


def add_body_paragraph(doc: Document, text: str) -> None:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text).replace("`", "")
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", text).replace("`", ""))


def add_quote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borderless(table)
    cell = table.cell(0, 0)
    shade(cell, PALE_CYAN)
    margins(cell, 150, 220, 150, 220)
    set_cell_text(cell, text.replace("**", "").replace("`", ""), color=CHARCOAL, size=11, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(doc: Document, lines: list[str], compact=False) -> None:
    rows = [[c.strip() for c in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.28] + [max(1.2, (6.75 - 1.28) / (len(rows[0]) - 1))] * (len(rows[0]) - 1) if len(rows[0]) > 1 else [6.75]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            shade(cell, INK if r_idx == 0 else (LIGHT if r_idx % 2 else WHITE))
            set_cell_text(cell, value.replace("**", "").replace("`", ""), color=WHITE if r_idx == 0 else INK, size=7.2 if compact else 8.2, bold=r_idx == 0 or c_idx == 0)
    table.rows[0].height = Inches(0.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_lifecycle(doc: Document) -> None:
    labels = ["01\nACTOR", "02\nIDENTITY", "03\nAUTHORITY", "04\nTRUST DECISION", "05\nCONTROL", "06\nEVIDENCE"]
    details = ["Human\nAI agent\nService", "Credential\nKey proof\nProvider", "Delegator\nScope\nExpiry", "Policy\nContext\nMemory", "ALLOW\nREVIEW\nDENY", "Receipt\nReplay\nMemory"]
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx in range(6):
        table.columns[idx].width = Inches(1.08)
        shade(table.cell(0, idx), CYAN if idx in (0, 5) else INK)
        shade(table.cell(1, idx), PALE_CYAN if idx in (0, 5) else LIGHT)
        set_cell_text(table.cell(0, idx), labels[idx], color=INK if idx in (0, 5) else WHITE, size=7.4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(1, idx), details[idx], size=7.3, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_authority_hero(doc: Document) -> None:
    add_quote(doc, "CONCEPTUAL OPERATIONAL REPRESENTATION")
    chain = [
        ("CHIEF FINANCIAL OFFICER", "Accountable delegator"),
        ("FINANCE AUTOMATION SERVICE", "Bounded service authority"),
        ("PAYMENTS AGENT", "Verified Ed25519 identity"),
        ("payment.create  ·  €14,500", "Treasury API · autonomous threshold €10,000"),
    ]
    for index, (name, detail) in enumerate(chain):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(2.35)
        table.columns[1].width = Inches(4.15)
        shade(table.cell(0, 0), INK)
        shade(table.cell(0, 1), LIGHT)
        set_cell_text(table.cell(0, 0), name, color=WHITE, size=8.2, bold=True)
        set_cell_text(table.cell(0, 1), detail, size=8.2)
        if index < len(chain) - 1:
            p = doc.add_paragraph("↓ delegates" if index < 2 else "↓ requests")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.runs[0].font.size = Pt(7.5)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor.from_string(SOFT_GREY)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(table.cell(0, 0), AMBER)
    shade(table.cell(0, 1), "FFF4D8")
    set_cell_text(table.cell(0, 0), "REVIEW", color=WHITE, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "Human approval required: amount exceeds delegated autonomous threshold.", size=9, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_decision_cards(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, color, copy) in enumerate((
        ("ALLOW", GREEN, "Authorized for this context"),
        ("REVIEW", AMBER, "Accountable intervention required"),
        ("DENY", RED, "Not authorized"),
    )):
        shade(table.cell(0, idx), color)
        set_cell_text(table.cell(0, idx), f"{label}\n{copy}", color=WHITE, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_receipt(doc: Document) -> None:
    rows = [
        ("TRANSACTION", "tx_…72d"), ("AGENT", "agent:payments"),
        ("ACTION", "payment.create"), ("RESOURCE", "treasury-api"),
        ("AUTHORITY", "grant v3 · current"), ("DECISION", "REVIEW"),
        ("REASON", "AMOUNT_EXCEEDS_AUTONOMOUS_THRESHOLD"),
        ("EVIDENCE", "identity · authority · policy · context"),
        ("REPLAY", "replay_…91f"),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.0)
    shade(table.cell(0, 0), INK); shade(table.cell(0, 1), INK)
    set_cell_text(table.cell(0, 0), "DECISION RECEIPT", color=CYAN, size=9, bold=True)
    set_cell_text(table.cell(0, 1), "Persisted · integrity-linked", color=WHITE, size=8, bold=True)
    for idx, (label, value) in enumerate(rows, 1):
        shade(table.cell(idx, 0), LIGHT)
        shade(table.cell(idx, 1), WHITE)
        set_cell_text(table.cell(idx, 0), label, size=7.5, bold=True)
        set_cell_text(table.cell(idx, 1), value, size=7.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_timeline(doc: Document) -> None:
    events = [
        ("10:21:03", "AGENT AUTHENTICATED"),
        ("10:21:04", "AUTHORITY RESOLVED"),
        ("10:21:04", "POLICY EVALUATED"),
        ("10:21:05", "REVIEW"),
        ("10:22:14", "HUMAN APPROVAL"),
        ("10:22:15", "FRESH EVALUATION · ALLOW"),
    ]
    table = doc.add_table(rows=len(events), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.25)
    table.columns[1].width = Inches(5.3)
    for idx, (timestamp, event) in enumerate(events):
        shade(table.cell(idx, 0), CYAN if idx in (0, len(events)-1) else INK)
        shade(table.cell(idx, 1), PALE_CYAN if idx in (0, len(events)-1) else LIGHT)
        set_cell_text(table.cell(idx, 0), timestamp, color=INK if idx in (0, len(events)-1) else WHITE, size=8, bold=True)
        set_cell_text(table.cell(idx, 1), event, size=8.5, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def blocks(markdown: str) -> list[tuple[str, object]]:
    lines = markdown.splitlines()
    result: list[tuple[str, object]] = []
    i = 0
    in_code = False
    code: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                result.append(("code", code)); code = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code.append(line); i += 1; continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i]); i += 1
            result.append(("table", table_lines)); continue
        if not line or line == "---":
            i += 1; continue
        if line.startswith("> "):
            result.append(("quote", line[2:])); i += 1; continue
        if line.startswith("- "):
            result.append(("bullet", line[2:])); i += 1; continue
        if re.match(r"^\d+\. ", line):
            result.append(("bullet", re.sub(r"^\d+\. ", "", line))); i += 1; continue
        if line.startswith("### "):
            result.append(("h3", line[4:])); i += 1; continue
        if line.startswith("## ") or line.startswith("# "):
            i += 1; continue
        paragraph = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,3} |[-|>] |\d+\. |```|---$)", lines[i]):
            paragraph.append(lines[i].strip()); i += 1
        result.append(("paragraph", " ".join(paragraph)))
    return result


def render_blocks(doc: Document, content: str, section_number: int | None = None) -> None:
    for kind, value in blocks(content):
        if kind == "paragraph": add_body_paragraph(doc, str(value))
        elif kind == "bullet": add_bullet(doc, str(value))
        elif kind == "quote": add_quote(doc, str(value))
        elif kind == "h3": doc.add_paragraph(str(value), style="Heading 2")
        elif kind == "table": add_markdown_table(doc, value, compact=section_number == 18)
        elif kind == "code" and section_number != 7:
            table = doc.add_table(rows=1, cols=1); borderless(table); shade(table.cell(0, 0), INK)
            set_cell_text(table.cell(0, 0), "\n".join(value), color=WHITE, size=7.4)


def main() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    exec_match = re.search(r"## Executive abstract\n(.*?)(?=\n---\n\n## 1\.)", source, re.S)
    section_matches = list(re.finditer(r"^## (\d+)\. (.+)$", source, re.M))
    if not exec_match or len(section_matches) != 20:
        raise SystemExit("Expected one executive abstract and exactly 20 numbered sections.")

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    doc.add_page_break()

    page_header(doc, "00", "Executive abstract", "WHY OPERATIONAL TRUST EXISTS")
    add_quote(doc, "WHO is acting?  ·  WHAT was delegated?  ·  IS authority still valid?  ·  SHOULD this action execute?  ·  CAN the enterprise prove why?")
    render_blocks(doc, exec_match.group(1).strip())

    for idx, match in enumerate(section_matches):
        number = int(match.group(1))
        title = match.group(2).strip()
        end = section_matches[idx + 1].start() if idx + 1 < len(section_matches) else source.find("\n---\n\n## Verification references", match.end())
        content = source[match.end():end if end != -1 else len(source)].strip()
        doc.add_page_break()
        page_header(doc, f"{number:02}", title)
        if number == 3: add_lifecycle(doc)
        if number == 7: add_authority_hero(doc)
        if number == 8: add_decision_cards(doc)
        if number == 11: add_receipt(doc)
        if number == 12: add_timeline(doc)
        render_blocks(doc, content, section_number=number)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    sys.exit(main())
